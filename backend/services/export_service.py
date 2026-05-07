"""
导出服务。
支持 Word（按用户模板格式）和 PDF（打包中文字体）导出。
"""
import io
import re
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from paths import FONTS_DIR, DEFAULT_FONT_PATH


REPORTLAB_FONT_NAME = "CJK"


def _register_cjk_font():
    try:
        font_path = DEFAULT_FONT_PATH
        if font_path.exists():
            pdfmetrics.registerFont(TTFont(REPORTLAB_FONT_NAME, str(font_path)))
            return True
        else:
            return False
    except Exception:
        return False


_register_cjk_font()


def _parse_markdown(text: str) -> list:
    lines = text.strip().split("\n")
    blocks = []
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("|"):
            table_rows.append(stripped)
            in_table = True
            continue
        else:
            if in_table and table_rows:
                blocks.append(("table", table_rows))
                table_rows = []
            in_table = False

        if not stripped:
            blocks.append(("blank", ""))
            continue

        if stripped.startswith("# "):
            blocks.append(("h1", stripped[2:].strip()))
        elif stripped.startswith("## "):
            blocks.append(("h2", stripped[3:].strip()))
        elif stripped.startswith("### "):
            blocks.append(("h3", stripped[4:].strip()))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append(("bullet", stripped[2:].strip()))
        elif re.match(r"^\d+\. ", stripped):
            blocks.append(("numbered", re.sub(r"^\d+\. ", "", stripped)))
        elif stripped.startswith("> "):
            blocks.append(("quote", stripped[2:].strip()))
        else:
            blocks.append(("paragraph", _inline_format(stripped)))

    if in_table and table_rows:
        blocks.append(("table", table_rows))

    return blocks


def _inline_format(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    text = re.sub(r"`(.+?)`", r"<tt>\1</tt>", text)
    return text


def _parse_table_row(row: str) -> list:
    cells = re.split(r"\s*\|\s*", row.strip())
    return [c for c in cells if c != ""]


def _render_markdown_to_docx(doc: Document, markdown_text: str):
    blocks = _parse_markdown(markdown_text)

    for block_type, content in blocks:
        if block_type == "blank":
            continue

        if block_type == "h1":
            p = doc.add_heading(content, level=1)
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.bold = True

        elif block_type == "h2":
            p = doc.add_heading(content, level=2)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.bold = True

        elif block_type == "h3":
            p = doc.add_heading(content, level=3)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True

        elif block_type == "bullet":
            p = doc.add_paragraph(content, style="List Bullet")

        elif block_type == "numbered":
            p = doc.add_paragraph(content, style="List Number")

        elif block_type == "quote":
            p = doc.add_paragraph(content)
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)

        elif block_type == "paragraph":
            p = doc.add_paragraph()
            _add_formatted_text(p, content)

        elif block_type == "table":
            rows = content
            if len(rows) < 2:
                continue
            header_row = _parse_table_row(rows[0])
            table = doc.add_table(rows=1, cols=len(header_row))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, cell_text in enumerate(header_row):
                hdr_cells[i].text = cell_text
                for para in hdr_cells[i].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

            for row_text in rows[2:]:
                cells = _parse_table_row(row_text)
                if cells:
                    row_cells = table.add_row().cells
                    for i, cell_text in enumerate(cells):
                        if i < len(row_cells):
                            row_cells[i].text = cell_text

            doc.add_paragraph()


def _add_formatted_text(para, text: str):
    parts = re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
    for part in parts:
        run = para.add_run(part)
        run.font.size = Pt(11)
        if part.startswith("**") and part.endswith("**"):
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.font.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.font.name = "Courier New"
            run.font.size = Pt(10)


def export_to_word(markdown_text: str, template_path: Optional[str] = None) -> bytes:
    if template_path:
        doc = Document(template_path)
        _render_markdown_to_docx(doc, markdown_text)
    else:
        doc = Document()
        doc.add_heading("工作报告", 0)
        _render_markdown_to_docx(doc, markdown_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_to_pdf(markdown_text: str) -> bytes:
    blocks = _parse_markdown(markdown_text)

    buffer = io.BytesIO()
    use_cjk = _register_cjk_font()

    if use_cjk:
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )
    else:
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

    styles = getSampleStyleSheet()

    if use_cjk:
        base_font = REPORTLAB_FONT_NAME
    else:
        base_font = "Helvetica"

    body_style = ParagraphStyle(
        "CJKBody",
        parent=styles["Normal"],
        fontName=base_font,
        fontSize=11,
        leading=16,
    )
    h1_style = ParagraphStyle(
        "CJKH1",
        parent=styles["Heading1"],
        fontName=base_font,
        fontSize=18,
        leading=24,
        spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        "CJKH2",
        parent=styles["Heading2"],
        fontName=base_font,
        fontSize=14,
        leading=20,
        spaceAfter=8,
    )
    h3_style = ParagraphStyle(
        "CJKH3",
        parent=styles["Heading3"],
        fontName=base_font,
        fontSize=12,
        leading=18,
        spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        "CJKBullet",
        parent=body_style,
        leftIndent=20,
        bulletIndent=10,
        spaceBefore=2,
        spaceAfter=2,
    )
    quote_style = ParagraphStyle(
        "CJKQuote",
        parent=body_style,
        leftIndent=30,
        textColor=(0.4, 0.4, 0.4),
        fontName=base_font,
        fontSize=10,
    )

    story = []
    for block_type, content in blocks:
        if block_type == "blank":
            story.append(Spacer(1, 4))
        elif block_type == "h1":
            story.append(Paragraph(content, h1_style))
        elif block_type == "h2":
            story.append(Paragraph(content, h2_style))
        elif block_type == "h3":
            story.append(Paragraph(content, h3_style))
        elif block_type == "bullet":
            story.append(Paragraph(f"• {content}", bullet_style))
        elif block_type == "numbered":
            story.append(Paragraph(content, body_style))
        elif block_type == "quote":
            story.append(Paragraph(content, quote_style))
        elif block_type == "paragraph":
            story.append(Paragraph(content, body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()
